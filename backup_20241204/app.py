from flask import Flask, render_template, request, jsonify, send_file, session
from flask_cors import CORS
from werkzeug.utils import secure_filename
import os
import json
from datetime import datetime, timedelta
from io import BytesIO
import anthropic
import time
from tenacity import retry, stop_after_attempt, wait_exponential
from PyPDF2 import PdfReader
import re
from dotenv import load_dotenv
import chardet
import pickle
from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from flask_session import Session
import logging
import logging.handlers
import sys

# Configure logging
def setup_logging():
    log_formatter = logging.Formatter(
        '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    
    # File handler for detailed logs
    log_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'logs', 'genas.log')
    os.makedirs(os.path.dirname(log_file), exist_ok=True)
    
    file_handler = logging.handlers.RotatingFileHandler(
        log_file,
        maxBytes=10*1024*1024,  # 10MB
        backupCount=5
    )
    file_handler.setFormatter(log_formatter)
    file_handler.setLevel(logging.DEBUG)
    
    # Console handler for standard output
    console_handler = logging.StreamHandler(sys.stdout)
    console_handler.setFormatter(log_formatter)
    console_handler.setLevel(logging.INFO)
    
    # Setup root logger
    root_logger = logging.getLogger()
    root_logger.setLevel(logging.DEBUG)
    root_logger.addHandler(file_handler)
    root_logger.addHandler(console_handler)
    
    # Setup specific logger for Anthropic calls
    anthropic_logger = logging.getLogger('anthropic')
    anthropic_logger.setLevel(logging.DEBUG)
    
    return anthropic_logger

# Initialize logging
anthropic_logger = setup_logging()

# Load environment variables
load_dotenv()

# Initialize Flask app
app = Flask(__name__)
app.secret_key = os.getenv('FLASK_SECRET_KEY', 'default-secret-key')  # Set a secret key for session management
CORS(app)  # Enable CORS for all routes

# Increase maximum content length to 20MB
app.config['MAX_CONTENT_LENGTH'] = 20 * 1024 * 1024  # 20MB in bytes
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
app.config['SESSION_TYPE'] = 'filesystem'
app.config['SESSION_FILE_DIR'] = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'flask_session')
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=24)
app.config['SESSION_PERMANENT'] = True

# Ensure session directory exists
os.makedirs(app.config['SESSION_FILE_DIR'], exist_ok=True)

# Initialize session interface
Session(app)

# Default prompt for AI interactions
DEFAULT_PROMPT = """You are an AI assistant specialized in DoD Acquisition Strategy documents. Your role is to:
1. Analyze the provided documents thoroughly
2. Answer questions based on the document content first, then supplement with your knowledge
3. Always cite which document you're referencing in your response
4. If relevant, suggest appropriate sections for including the information in an Acquisition Strategy

Please provide clear, concise responses that help create or improve DoD Acquisition Strategy documents."""

# System prompt for all interactions
SYSTEM_PROMPT = """You are an AI assistant specialized in DoD Acquisition Strategy documents. Your role is to:
1. Analyze the provided documents thoroughly
2. Answer questions based on the document content first, then supplement with your knowledge
3. Always cite which document you're referencing in your response
4. If relevant, suggest appropriate sections for including the information in an Acquisition Strategy"""

app.config['VERSIONS_FOLDER'] = 'versions'
app.config['DOCUMENTS_STORE'] = 'documents_store.pkl'

# Ensure required directories exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['VERSIONS_FOLDER'], exist_ok=True)

# Load existing document contents from disk
try:
    with open(app.config['DOCUMENTS_STORE'], 'rb') as f:
        document_contents = pickle.load(f)
    print(f"Loaded {len(document_contents)} documents from store")
except FileNotFoundError:
    document_contents = {}
    print("No existing document store found, starting fresh")
except Exception as e:
    print(f"Error loading document store: {e}")
    document_contents = {}

def save_documents_store():
    """Save the documents store to disk"""
    try:
        with open(app.config['DOCUMENTS_STORE'], 'wb') as f:
            pickle.dump(document_contents, f)
        print(f"Saved {len(document_contents)} documents to store")
    except Exception as e:
        print(f"Error saving document store: {e}")

# Store uploaded documents content
document_contents = {}

def extract_text_from_pdf(file_path):
    try:
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            text = ''
            for page in reader.pages:
                text += page.extract_text() + '\n'
            return text
    except Exception as e:
        print(f"Error extracting PDF text: {str(e)}")
        return None

def extract_text_from_docx(file_path):
    """Extract text from a .docx file"""
    try:
        doc = Document(file_path)
        text = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
        
        return '\n'.join(text)
    except Exception as e:
        print(f"Error extracting DOCX text: {str(e)}")
        return None

def detect_encoding(file_path):
    with open(file_path, 'rb') as file:
        raw_data = file.read()
        result = chardet.detect(raw_data)
        return result['encoding']

def extract_text_from_file(file_path):
    """Extract text from various file types"""
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    else:
        try:
            # Try to detect the file encoding
            encoding = detect_encoding(file_path)
            if not encoding:
                encoding = 'utf-8'
            
            with open(file_path, 'r', encoding=encoding, errors='replace') as file:
                return file.read()
        except Exception as e:
            print(f"Error reading file: {str(e)}")
            try:
                with open(file_path, 'rb') as file:
                    content = file.read()
                    return content.decode('utf-8', errors='replace')
            except Exception as e:
                print(f"Error in fallback reading: {str(e)}")
                return None

# Define the sections of the Acquisition Strategy document
ACQUISITION_SECTIONS = [
    "Executive Summary",
    "Program Overview",
    "Risk Assessment",
    "Market Research",
    "Competition Strategy",
    "Source Selection Planning",
    "Business Considerations",
    "Multi-Year Procurement",
    "Lease-Purchase Analysis",
    "Source of Support",
    "Environmental Considerations",
    "Security Considerations",
    "Make or Buy Program",
    "Contract Types",
    "Sustainment Strategy",
    "Supporting Documents",
    "Scratch Pad"
]

@retry(
    stop=stop_after_attempt(int(os.getenv('ANTHROPIC_MAX_RETRIES', 3))),
    wait=wait_exponential(
        multiplier=float(os.getenv('ANTHROPIC_BACKOFF_FACTOR', 2)),
        min=4,
        max=60
    ),
    retry_error_callback=lambda retry_state: {'error': f'API temporarily unavailable after {retry_state.attempt_number} attempts. Please try again in a few minutes.'}
)
def get_claude_response(system_prompt, user_prompt):
    """Get response from Claude API with retry logic"""
    try:
        anthropic_logger.info("\n=== Starting Claude API request ===")
        anthropic_logger.debug(f"System prompt: {system_prompt}")
        anthropic_logger.debug(f"User prompt: {user_prompt}")
        anthropic_logger.info(f"System prompt length: {len(system_prompt)}")
        anthropic_logger.info(f"User prompt length: {len(user_prompt)}")
        anthropic_logger.info(f"Retry settings: max_retries={os.getenv('ANTHROPIC_MAX_RETRIES', 3)}, backoff_factor={os.getenv('ANTHROPIC_BACKOFF_FACTOR', 2)}")
        
        # Debug: Print API key status (safely)
        api_key = os.getenv('ANTHROPIC_API_KEY')
        if not api_key:
            anthropic_logger.error("ERROR: No API key found!")
            raise Exception("Anthropic API key not found in environment variables")
        else:
            anthropic_logger.info(f"API key verification: {api_key[:8]}... (length: {len(api_key)})")
            
        try:
            anthropic_logger.info("\nInitializing Anthropic client...")
            client = anthropic.Client(api_key=api_key)
            
            anthropic_logger.info("Preparing message for Claude...")
            model_name = os.getenv('MODEL_NAME', 'claude-2.1')
            anthropic_logger.info(f"Using model: {model_name}")
            anthropic_logger.info(f"Total prompt length: {len(system_prompt) + len(user_prompt)}")
            
            anthropic_logger.info("\nSending request to Claude API...")
            try:
                message = client.messages.create(
                    model=model_name,
                    max_tokens=int(os.getenv('MAX_TOKENS', 2000)),
                    temperature=float(os.getenv('TEMPERATURE', 0.7)),
                    system=system_prompt,
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "text",
                                    "text": user_prompt
                                }
                            ]
                        }
                    ]
                )
                
                anthropic_logger.debug(f"Claude response: {message}")
                anthropic_logger.info("=== Claude API request completed successfully ===")
                return message
                
            except anthropic.APIError as api_error:
                anthropic_logger.error(f"\nAnthropicAPI Error Details:")
                anthropic_logger.error(f"Error Status: {api_error.status_code if hasattr(api_error, 'status_code') else 'Unknown'}")
                anthropic_logger.error(f"Error Type: {api_error.type if hasattr(api_error, 'type') else 'Unknown'}")
                anthropic_logger.error(f"Error Message: {str(api_error)}")
                anthropic_logger.error(f"Full Error: {api_error.__dict__}")
                raise
            except Exception as e:
                anthropic_logger.error(f"\nUnexpected Error in API call:")
                anthropic_logger.error(f"Error Type: {type(e)}")
                anthropic_logger.error(f"Error Message: {str(e)}")
                anthropic_logger.error(f"Error Details: {e.__dict__ if hasattr(e, '__dict__') else 'No details available'}")
                raise
        except Exception as e:
            anthropic_logger.error("\n!!! General Error in get_claude_response !!!")
            anthropic_logger.error(f"Error message: {str(e)}")
            anthropic_logger.error(f"Error type: {type(e)}")
            raise
    except Exception as e:
        anthropic_logger.error(f"Final error in get_claude_response: {str(e)}")
        return {'error': str(e)}

@app.route('/')
def index():
    return render_template('index.html', sections=ACQUISITION_SECTIONS)

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
        
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
        
    if file:
        try:
            # Check file size before processing
            file.seek(0, os.SEEK_END)
            size = file.tell()
            file.seek(0)
            
            if size > app.config['MAX_CONTENT_LENGTH']:
                return jsonify({
                    'error': f'File too large. Maximum size is {app.config["MAX_CONTENT_LENGTH"] // (1024 * 1024)}MB'
                }), 413

            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            
            # Create upload directory if it doesn't exist
            os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
            
            # Save the file
            file.save(file_path)
            
            print(f"Processing file: {filename} (size: {size / 1024:.1f}KB)")
            
            # Extract and store document content
            content = extract_text_from_file(file_path)
            if content is None:
                os.remove(file_path)  # Clean up the file if we can't process it
                return jsonify({'error': 'Could not extract text from file'}), 400
                
            # Store in session
            if 'uploaded_documents' not in session:
                session['uploaded_documents'] = {}
            
            # Check content size
            content_size = len(content.encode('utf-8'))
            if content_size > app.config['MAX_CONTENT_LENGTH']:
                os.remove(file_path)
                return jsonify({
                    'error': f'Extracted content too large. Maximum size is {app.config["MAX_CONTENT_LENGTH"] // (1024 * 1024)}MB'
                }), 413
            
            session['uploaded_documents'][filename] = content
            session.modified = True
            
            print(f"Successfully processed {filename}. Content size: {content_size / 1024:.1f}KB")
            
            return jsonify({
                'message': 'File uploaded and processed successfully',
                'filename': filename,
                'contentSize': content_size
            })
            
        except Exception as e:
            print(f"Upload error: {str(e)}")
            # Clean up file if it exists
            if 'file_path' in locals() and os.path.exists(file_path):
                os.remove(file_path)
            return jsonify({'error': f'Error processing file: {str(e)}'}), 500
            
    return jsonify({'error': 'Invalid file'}), 400

@app.route('/save_version', methods=['POST'])
def save_version():
    """Save current document state as a version"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400

        version_name = data.get('version')
        sections = data.get('sections', [])
        
        if not version_name:
            return jsonify({'error': 'Version name is required'}), 400
            
        # Ensure version name has .json extension
        if not version_name.endswith('.json'):
            version_name += '.json'
            
        # Create versions directory if it doesn't exist
        versions_dir = app.config['VERSIONS_FOLDER']
        os.makedirs(versions_dir, exist_ok=True)
        
        version_path = os.path.join(versions_dir, secure_filename(version_name))
        
        # Save version data
        version_data = {
            'version': version_name,
            'timestamp': datetime.now().isoformat(),
            'sections': sections
        }
        
        app.logger.info(f'Saving version data: {version_data}')  # Debug log
        
        with open(version_path, 'w') as f:
            json.dump(version_data, f, indent=2)
        
        return jsonify({'message': 'Version saved successfully', 'version': version_name})
    except Exception as e:
        app.logger.error(f'Error saving version: {str(e)}')
        return jsonify({'error': str(e)}), 500

@app.route('/load_version/<version>', methods=['GET'])
def load_version(version):
    """Load a specific version"""
    try:
        version_path = os.path.join(app.config['VERSIONS_FOLDER'], secure_filename(version))
        if not os.path.exists(version_path):
            return jsonify({'error': 'Version not found'}), 404
            
        with open(version_path, 'r') as f:
            version_data = json.load(f)
            
        # Add sections array if it doesn't exist (for backward compatibility)
        if 'sections' not in version_data:
            app.logger.warning(f'Converting old version format for {version}')
            # Try to handle old format
            if 'content' in version_data:
                sections = []
                for section_data in version_data['content']:
                    sections.append({
                        'title': section_data.get('title', 'Unknown Section'),
                        'content': section_data.get('content', '')
                    })
                version_data['sections'] = sections
            else:
                return jsonify({'error': 'Invalid version file format'}), 400
            
        app.logger.info(f'Loading version data: {version_data}')  # Debug log
        return jsonify(version_data)
    except Exception as e:
        app.logger.error(f'Error loading version {version}: {str(e)}')
        return jsonify({'error': str(e)}), 500

@app.route('/delete_version/<version>', methods=['DELETE'])
def delete_version(version):
    """Delete a specific version"""
    try:
        version_path = os.path.join(app.config['VERSIONS_FOLDER'], secure_filename(version))
        if not os.path.exists(version_path):
            return jsonify({'error': 'Version not found'}), 404
            
        os.remove(version_path)
        app.logger.info(f'Deleted version: {version}')
        return jsonify({'message': f'Version {version} deleted successfully'})
    except Exception as e:
        app.logger.error(f'Error deleting version {version}: {str(e)}')
        return jsonify({'error': str(e)}), 500

@app.route('/list_versions', methods=['GET'])
def list_versions():
    """List all saved versions"""
    try:
        versions_dir = app.config['VERSIONS_FOLDER']
        if not os.path.exists(versions_dir):
            return jsonify({'versions': []})
            
        versions = []
        for filename in os.listdir(versions_dir):
            if filename.endswith('.json'):
                file_path = os.path.join(versions_dir, filename)
                try:
                    with open(file_path, 'r') as f:
                        data = json.load(f)
                        # Ensure the version has the correct format
                        if 'sections' in data or 'content' in data:
                            versions.append(filename)
                except Exception as e:
                    app.logger.warning(f'Error reading version file {filename}: {str(e)}')
                    continue
        
        versions.sort(key=lambda x: os.path.getmtime(os.path.join(versions_dir, x)), reverse=True)
        return jsonify({'versions': versions})
    except Exception as e:
        app.logger.error(f'Error listing versions: {str(e)}')
        return jsonify({'error': str(e)}), 500

@app.route('/get_versions', methods=['GET'])
def get_versions():
    versions = []
    if os.path.exists(app.config['VERSIONS_FOLDER']):
        for filename in os.listdir(app.config['VERSIONS_FOLDER']):
            if filename.endswith('.json'):
                file_path = os.path.join(app.config['VERSIONS_FOLDER'], filename)
                with open(file_path) as f:
                    version_data = json.load(f)
                    versions.append({
                        'name': version_data['version'],
                        'timestamp': version_data['timestamp']
                    })
    
    return jsonify({'versions': sorted(versions, key=lambda x: x['timestamp'], reverse=True)})

@app.route('/get_version/<version_name>', methods=['GET'])
def get_version(version_name):
    version_path = os.path.join(app.config['VERSIONS_FOLDER'], f'{version_name}.json')
    if os.path.exists(version_path):
        with open(version_path) as f:
            return jsonify(json.load(f))
    return jsonify({'error': 'Version not found'}), 404

@app.route('/documents', methods=['GET'])
def list_documents():
    """List all uploaded documents"""
    return jsonify({
        'documents': list(document_contents.keys()),
        'count': len(document_contents)
    })

@app.route('/delete_document/<filename>', methods=['DELETE'])
def delete_document(filename):
    """Delete a document from the store"""
    try:
        if filename in document_contents:
            # Remove from document store
            del document_contents[filename]
            save_documents_store()
            
            # Remove physical file if it exists
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            if os.path.exists(file_path):
                os.remove(file_path)
            
            return jsonify({'message': f'Document {filename} deleted successfully'})
        return jsonify({'error': 'Document not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/get_required_documents', methods=['POST'])
def get_required_documents():
    """Get required supporting documents based on current content"""
    try:
        # Return a simple list of required documents for now
        # This can be enhanced with AI-based document suggestions later
        required_docs = [
            {"name": "Systems Engineering Plan (SEP)", "applicability": 85},
            {"name": "Test and Evaluation Master Plan (TEMP)", "applicability": 75},
            {"name": "Cybersecurity Strategy", "applicability": 90},
            {"name": "Life Cycle Sustainment Plan (LCSP)", "applicability": 80}
        ]
        
        return jsonify({"documents": required_docs})
    except Exception as e:
        print(f"Error getting required documents: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/chat', methods=['POST'])
def chat():
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data received'}), 400

        message = data.get('message', '')
        include_sections = data.get('includeSections', False)
        current_content = data.get('currentContent', {})
        include_documents = data.get('includeDocuments', True)
        uploaded_documents = data.get('uploadedDocuments', [])
        prompt_id = data.get('promptId', '')

        print('Received chat request:', {
            'messageLength': len(message),
            'includeSections': include_sections,
            'includeDocuments': include_documents,
            'uploadedDocs': uploaded_documents,
            'numSections': len(current_content),
            'sectionNames': list(current_content.keys()),
            'promptId': prompt_id
        })

        if not message:
            return jsonify({'error': 'No message provided'}), 400

        # Get target section from prompt if available
        target_section = None
        if prompt_id:
            prompts = load_prompts()
            for prompt in prompts.get('prompts', []):
                if prompt.get('id') == prompt_id:
                    target_section = prompt.get('targetSection')
                    break
            print('Found target section from prompt:', target_section)

        # Build system prompt
        system_prompt = DEFAULT_PROMPT

        # Add document context
        document_context = []
        if include_documents and 'uploaded_documents' in session:
            uploaded_docs = session.get('uploaded_documents', {})
            print(f"Found {len(uploaded_docs)} documents in session")
            for doc_name, content in uploaded_docs.items():
                if not uploaded_documents or doc_name in uploaded_documents:
                    document_context.append(f"Document: {doc_name}\nContent: {content}\n")

        # Add section context
        section_context = []
        if include_sections and current_content:
            for section, content in current_content.items():
                if content and content.strip():
                    section_context.append(f"Section: {section}\nContent: {content}\n")

        # Combine all context
        user_prompt = "I have provided the following documents and sections for reference:\n\n"
        
        if document_context:
            user_prompt += "=== Documents ===\n" + "\n".join(document_context) + "\n\n"
        else:
            print("Warning: No document context available")
            
        if section_context:
            user_prompt += "=== Sections ===\n" + "\n".join(section_context) + "\n\n"
            
        user_prompt += f"Question: {message}\n\n"
        
        if not document_context and not section_context:
            return jsonify({
                'error': 'No documents or sections available. Please upload documents or provide section content first.'
            }), 400

        print(f"Generated prompt with {len(document_context)} documents and {len(section_context)} sections")

        # Get response from Claude
        try:
            response = get_claude_response(system_prompt, user_prompt)
            
            if isinstance(response, dict) and 'error' in response:
                return jsonify(response), 500

            # Extract text from Claude-3 Message object
            try:
                # Access the text content directly from the Message object
                if hasattr(response, 'content') and isinstance(response.content, list) and len(response.content) > 0:
                    # Get the first content block's text
                    first_block = response.content[0]
                    if hasattr(first_block, 'text'):
                        response_text = first_block.text
                    else:
                        response_text = str(first_block)
                else:
                    response_text = str(response)
            except Exception as e:
                print(f"Error extracting text from response: {str(e)}")
                response_text = str(response)

            # Create a simple dict with only the text
            return jsonify({
                'response': response_text,
                'suggestedSection': None,
                'targetSection': target_section
            })

        except Exception as e:
            error_msg = f"Error processing Claude response: {str(e)}"
            print(error_msg)
            return jsonify({'error': error_msg}), 500

    except Exception as e:
        error_msg = f"Error in chat endpoint: {str(e)}"
        print(error_msg)
        print(f"Error type: {type(e)}")
        print(f"Error details: {e.__dict__ if hasattr(e, '__dict__') else 'No additional details'}")
        return jsonify({'error': error_msg}), 500

@app.route('/print_document', methods=['POST'])
def print_document():
    try:
        # Get the document content from the request
        data = request.get_json()
        if not data or 'sections' not in data:
            return jsonify({'error': 'No document content provided'}), 400

        sections = data['sections']
        
        # Create a BytesIO buffer for the PDF
        buffer = BytesIO()
        
        # Create the PDF document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )

        # Create styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=20,
            spaceAfter=10,
            textColor=colors.HexColor('#2c3e50')
        )
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceBefore=6,
            spaceAfter=12,
            leading=14
        )
        date_style = ParagraphStyle(
            'CustomDate',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.gray,
            alignment=1  # Center alignment
        )

        # Build the document content
        story = []
        
        # Add title
        story.append(Paragraph("Acquisition Strategy Document", title_style))
        story.append(Paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y')}", date_style))
        story.append(Spacer(1, 30))

        # Add sections
        for title, content in sections.items():
            if content.strip():
                story.append(Paragraph(title, heading_style))
                # Split content into paragraphs
                paragraphs = content.split('\n')
                for para in paragraphs:
                    if para.strip():
                        story.append(Paragraph(para.strip(), normal_style))
                story.append(Spacer(1, 12))

        # Build the PDF
        doc.build(story)

        # Reset buffer position
        buffer.seek(0)
        
        # Create response
        response = send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name='acquisition_strategy.pdf'
        )

        return response

    except Exception as e:
        print(f"Error generating PDF: {str(e)}")
        return jsonify({'error': 'Error generating PDF'}), 500

def get_or_create_document():
    doc_path = os.path.join(app.config['UPLOAD_FOLDER'], 'current_document.json')
    
    if not os.path.exists(doc_path):
        # Create default document structure
        default_document = {
            "title": "Acquisition Strategy",
            "sections": [
                {
                    "title": section,
                    "content": "",
                    "subsections": []
                }
                for section in ACQUISITION_SECTIONS
            ]
        }
        
        # Ensure upload directory exists
        os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
        
        # Save default document
        with open(doc_path, 'w') as f:
            json.dump(default_document, f, indent=4)
        
        return default_document
    
    # Load existing document
    with open(doc_path, 'r') as f:
        return json.load(f)

@app.route('/update_section', methods=['POST'])
def update_section():
    try:
        data = request.get_json()
        section = data.get('section')
        content = data.get('content')
        
        if not section or not content:
            return jsonify({'error': 'Missing section or content'}), 400
            
        # Get or create document
        document = get_or_create_document()
            
        # Update the specified section
        section_found = False
        for section_data in document['sections']:
            if section_data['title'] == section:
                section_data['content'] = content
                section_found = True
                break
                
        if not section_found:
            return jsonify({'error': f'Section {section} not found'}), 404
            
        # Save the updated document
        doc_path = os.path.join(app.config['UPLOAD_FOLDER'], 'current_document.json')
        with open(doc_path, 'w') as f:
            json.dump(document, f, indent=4)
            
        return jsonify({'message': f'Successfully updated {section}'}), 200
        
    except Exception as e:
        app.logger.error(f'Error updating section: {str(e)}')
        return jsonify({'error': 'Internal server error'}), 500

@app.route('/get_document', methods=['GET'])
def get_document():
    try:
        document = get_or_create_document()
        return jsonify(document)
    except Exception as e:
        app.logger.error(f'Error getting document: {str(e)}')
        return jsonify({'error': 'Internal server error'}), 500

@app.route('/clear_session', methods=['POST'])
def clear_session():
    """Clear the current session state"""
    try:
        print("Clearing session state...")
        
        # Clear document contents
        global document_contents
        document_contents = {}
        save_documents_store()
        
        # Clear uploaded files
        upload_dir = app.config['UPLOAD_FOLDER']
        for filename in os.listdir(upload_dir):
            file_path = os.path.join(upload_dir, filename)
            try:
                if os.path.isfile(file_path):
                    os.unlink(file_path)
                    print(f"Deleted file: {filename}")
            except Exception as e:
                print(f"Error deleting {filename}: {e}")
        
        print("Session cleared successfully")
        return jsonify({'message': 'Session cleared successfully'})
        
    except Exception as e:
        print(f"Error clearing session: {e}")
        return jsonify({'error': str(e)}), 500

def load_prompts():
    """Load prompts from JSON file"""
    try:
        with open('static/prompts.json', 'r') as f:
            prompts = json.load(f)
            print(f"Loaded {len(prompts)} prompts from prompts.json")
            return prompts
    except Exception as e:
        print(f"Error loading prompts: {e}")
        return {
            'default': {
                'id': 'default',
                'name': 'Default Prompt',
                'description': 'Default prompt for DoD Acquisition Strategy assistance',
                'prompt': DEFAULT_PROMPT
            }
        }

@app.route('/prompts', methods=['GET'])
def get_prompts():
    """Return all available prompts"""
    try:
        prompts = load_prompts()
        return jsonify(prompts)
    except Exception as e:
        print(f"Error in get_prompts: {e}")
        return jsonify({
            'prompts': [{
                'id': 'default',
                'name': 'Default Prompt',
                'description': 'Default prompt for DoD Acquisition Strategy assistance',
                'prompt': DEFAULT_PROMPT,
                'category': 'General'
            }]
        })

@app.route('/prompt/<prompt_id>', methods=['GET'])
def get_prompt(prompt_id):
    """Return a specific prompt by ID"""
    try:
        prompts = load_prompts()
        if 'prompts' in prompts:
            for prompt in prompts['prompts']:
                if prompt['id'] == prompt_id:
                    return jsonify(prompt)
        return jsonify({'error': 'Prompt not found'}), 404
    except Exception as e:
        print(f"Error in get_prompt: {e}")
        if prompt_id == 'default':
            return jsonify({
                'id': 'default',
                'name': 'Default Prompt',
                'description': 'Default prompt for DoD Acquisition Strategy assistance',
                'prompt': DEFAULT_PROMPT
            })
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    # Only use debug mode when running locally
    is_dev = os.getenv('FLASK_ENV') == 'development'
    port = int(os.getenv('PORT', 8000))
    app.run(host='0.0.0.0', port=port, debug=is_dev)
