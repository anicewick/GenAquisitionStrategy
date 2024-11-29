from flask import Flask, render_template, request, jsonify, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename
import os
import json
from datetime import datetime
from io import BytesIO
import anthropic
import time
from tenacity import retry, stop_after_attempt, wait_exponential
from PyPDF2 import PdfReader
import re
from dotenv import load_dotenv
import chardet
import pickle
from docx import Document

# Load environment variables
load_dotenv()

app = Flask(__name__)
CORS(app)  # Enable CORS for all routes
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['VERSIONS_FOLDER'] = 'versions'
app.config['DOCUMENTS_STORE'] = 'documents_store.pkl'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Initialize Anthropic client
client = anthropic.Anthropic(
    api_key=os.getenv('ANTHROPIC_API_KEY')
)

# Ensure required directories exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['VERSIONS_FOLDER'], exist_ok=True)

# Load existing document contents from disk
try:
    with open(app.config['DOCUMENTS_STORE'], 'rb') as f:
        document_contents = pickle.load(f)
    print(f"Loaded {len(document_contents)} documents from store")
except FileNotFoundError:
    document_contents = {}
    print("No existing document store found, starting fresh")
except Exception as e:
    print(f"Error loading document store: {e}")
    document_contents = {}

def save_documents_store():
    """Save the documents store to disk"""
    try:
        with open(app.config['DOCUMENTS_STORE'], 'wb') as f:
            pickle.dump(document_contents, f)
        print(f"Saved {len(document_contents)} documents to store")
    except Exception as e:
        print(f"Error saving document store: {e}")

# Store uploaded documents content
document_contents = {}

def extract_text_from_pdf(file_path):
    try:
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            text = ''
            for page in reader.pages:
                text += page.extract_text() + '\n'
            return text
    except Exception as e:
        print(f"Error extracting PDF text: {str(e)}")
        return None

def extract_text_from_docx(file_path):
    """Extract text from a .docx file"""
    try:
        doc = Document(file_path)
        text = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
        
        return '\n'.join(text)
    except Exception as e:
        print(f"Error extracting DOCX text: {str(e)}")
        return None

def detect_encoding(file_path):
    with open(file_path, 'rb') as file:
        raw_data = file.read()
        result = chardet.detect(raw_data)
        return result['encoding']

def extract_text_from_file(file_path):
    """Extract text from various file types"""
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    else:
        try:
            # Try to detect the file encoding
            encoding = detect_encoding(file_path)
            if not encoding:
                encoding = 'utf-8'
            
            with open(file_path, 'r', encoding=encoding, errors='replace') as file:
                return file.read()
        except Exception as e:
            print(f"Error reading file: {str(e)}")
            try:
                with open(file_path, 'rb') as file:
                    content = file.read()
                    return content.decode('utf-8', errors='replace')
            except Exception as e:
                print(f"Error in fallback reading: {str(e)}")
                return None

# Define the sections of the Acquisition Strategy document
ACQUISITION_SECTIONS = [
    "Purpose and Program Description",
    "Capability Need",
    "Market Research",
    "Business Considerations",
    "Risk",
    "Competition",
    "Incentives",
    "Product Support",
    "Joint and International",
    "Small Business Strategy",
    "Contracting Approach",
    "Test and Evaluation",
    "Cost, Funding and Schedule",
    "Approvals"
]

@retry(
    stop=stop_after_attempt(3),
    wait=wait_exponential(multiplier=1, min=4, max=10),
    retry_error_callback=lambda retry_state: {'error': f'API temporarily unavailable after {retry_state.attempt_number} attempts. Please try again in a few minutes.'}
)
def get_claude_response(system_prompt, user_prompt):
    """Get response from Claude API with retry logic"""
    try:
        print("Sending request to Claude API...")  # Debug log
        print(f"System prompt length: {len(system_prompt)}")
        print(f"User prompt length: {len(user_prompt)}")
        
        # Debug: Print API key status (safely)
        api_key = os.getenv('ANTHROPIC_API_KEY')
        if not api_key:
            print("WARNING: No API key found!")
        else:
            print(f"API key found (starts with: {api_key[:4]}...)")
            
        try:
            message = client.messages.create(
                model="claude-3-opus-20240229",
                max_tokens=2000,
                temperature=0.7,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": f"{system_prompt}\n\n{user_prompt}"
                            }
                        ]
                    }
                ]
            )
            
            if not message or not message.content:
                raise Exception("Empty response from Claude API")
                
            response_text = message.content[0].text
            print(f"Received response from Claude (length: {len(response_text)})")  # Debug log
            return response_text
            
        except anthropic.APIError as api_e:
            print(f"Detailed Claude API Error: {str(api_e)}")
            print(f"Error type: {type(api_e)}")
            print(f"Error attributes: {dir(api_e)}")
            if hasattr(api_e, 'status_code'):
                print(f"Status code: {api_e.status_code}")
            if "rate_limit" in str(api_e).lower():
                raise Exception("Rate limit exceeded. Please try again in a few minutes.")
            raise Exception(f"API Error: {str(api_e)}")
            
    except Exception as e:
        print(f"Detailed error in get_claude_response: {str(e)}")
        print(f"Error type: {type(e)}")
        print(f"Error attributes: {dir(e)}")
        raise

@app.route('/')
def index():
    return render_template('index.html', sections=ACQUISITION_SECTIONS)

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    if file:
        try:
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)
            
            # Extract and store document content
            content = extract_text_from_file(file_path)
            if content is None:
                return jsonify({'error': 'Could not extract text from file'}), 400
                
            document_contents[filename] = content
            save_documents_store()  # Save to disk after successful upload
            
            return jsonify({
                'message': 'File uploaded and processed successfully',
                'filename': filename
            })
        except Exception as e:
            print(f"Upload error: {str(e)}")
            return jsonify({'error': f'Error processing file: {str(e)}'}), 500

@app.route('/save_version', methods=['POST'])
def save_version():
    data = request.json
    content = data.get('content', {})
    version_name = data.get('version_name', f'version_{datetime.now().strftime("%Y%m%d_%H%M%S")}')
    
    version_path = os.path.join(app.config['VERSIONS_FOLDER'], f'{version_name}.json')
    version_data = {
        'content': content,
        'timestamp': datetime.now().isoformat(),
        'version_name': version_name
    }
    
    with open(version_path, 'w') as f:
        json.dump(version_data, f, indent=2)
    
    return jsonify({'message': 'Version saved successfully', 'version_name': version_name})

@app.route('/get_versions', methods=['GET'])
def get_versions():
    versions = []
    if os.path.exists(app.config['VERSIONS_FOLDER']):
        for filename in os.listdir(app.config['VERSIONS_FOLDER']):
            if filename.endswith('.json'):
                file_path = os.path.join(app.config['VERSIONS_FOLDER'], filename)
                with open(file_path) as f:
                    version_data = json.load(f)
                    versions.append({
                        'name': version_data['version_name'],
                        'timestamp': version_data['timestamp']
                    })
    
    return jsonify({'versions': sorted(versions, key=lambda x: x['timestamp'], reverse=True)})

@app.route('/get_version/<version_name>', methods=['GET'])
def get_version(version_name):
    version_path = os.path.join(app.config['VERSIONS_FOLDER'], f'{version_name}.json')
    if os.path.exists(version_path):
        with open(version_path) as f:
            return jsonify(json.load(f))
    return jsonify({'error': 'Version not found'}), 404

@app.route('/documents', methods=['GET'])
def list_documents():
    """List all uploaded documents"""
    return jsonify({
        'documents': list(document_contents.keys()),
        'count': len(document_contents)
    })

@app.route('/documents/<filename>', methods=['DELETE'])
def delete_document(filename):
    """Delete a document from the store"""
    if filename in document_contents:
        del document_contents[filename]
        save_documents_store()
        return jsonify({'message': f'Document {filename} deleted successfully'})
    return jsonify({'error': 'Document not found'}), 404

@app.route('/get_required_documents', methods=['POST'])
def get_required_documents():
    """Get required supporting documents based on current content"""
    try:
        # Return a simple list of required documents for now
        # This can be enhanced with AI-based document suggestions later
        required_docs = [
            {"name": "Systems Engineering Plan (SEP)", "applicability": 85},
            {"name": "Test and Evaluation Master Plan (TEMP)", "applicability": 75},
            {"name": "Cybersecurity Strategy", "applicability": 90},
            {"name": "Life Cycle Sustainment Plan (LCSP)", "applicability": 80}
        ]
        
        return jsonify({"documents": required_docs})
    except Exception as e:
        print(f"Error getting required documents: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/chat', methods=['POST'])
def chat():
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No JSON data received'}), 400
            
        message = data.get('message')
        if not message:
            return jsonify({'error': 'No message provided'}), 400
            
        print(f"Received chat message: {message}")  # Debug log
        
        if not document_contents:
            return jsonify({
                'response': 'No documents have been uploaded yet. Please upload a document first.',
                'suggestedSection': None,
                'canPropagate': False
            })
        
        # Prepare context from uploaded documents
        context = "\n\n".join([
            f"[Document: {filename}]\n{content}\n[End of Document: {filename}]"
            for filename, content in document_contents.items()
        ])
        
        print(f"Context prepared from {len(document_contents)} documents")  # Debug log
        
        # Prepare the prompt with clear instructions
        system_prompt = """You are an AI assistant specialized in DoD Acquisition Strategy documents. Your role is to:
1. Analyze the provided documents thoroughly
2. Answer questions based on the document content first, then supplement with your knowledge
3. Always cite which document you're referencing in your response
4. If relevant, suggest appropriate sections for including the information in an Acquisition Strategy"""

        user_prompt = f"""I have provided the following documents for reference:

{context}

Based on these documents, please answer the following question:
{message}

Remember to:
1. Reference specific documents in your response
2. Quote relevant passages when appropriate
3. If the information would fit in an Acquisition Strategy, suggest the relevant section"""
        
        try:
            print("Calling Claude API...")  # Debug log
            ai_response = get_claude_response(system_prompt, user_prompt)
            print("Received response from Claude")  # Debug log
            
            if not ai_response:
                raise Exception("Empty response from Claude")
            
            # Try to identify suggested section
            suggested_section = None
            response_text = str(ai_response)  # Ensure we have a string
            for section in ACQUISITION_SECTIONS:
                if section.lower() in response_text.lower():
                    suggested_section = section
                    break
            
            response_data = {
                'response': response_text,
                'suggestedSection': suggested_section,
                'canPropagate': suggested_section is not None
            }
            
            print("Sending response to client")  # Debug log
            return jsonify(response_data)
            
        except Exception as e:
            error_message = str(e)
            print(f"Error in Claude API call: {error_message}")  # Debug log
            return jsonify({
                'error': 'Failed to get AI response. Please try again.',
                'details': error_message
            }), 500
        
    except Exception as e:
        error_message = str(e)
        print(f"Error in chat endpoint: {error_message}")  # Debug log
        return jsonify({
            'error': 'Server error occurred',
            'details': error_message
        }), 500

if __name__ == '__main__':
    app.run(debug=True)
