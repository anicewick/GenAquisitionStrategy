from flask import Flask, render_template, request, jsonify, send_file, session
from flask_cors import CORS
from werkzeug.utils import secure_filename
import os
import json
from datetime import datetime, timedelta
from io import BytesIO
import anthropic
import time
from tenacity import retry, stop_after_attempt, wait_exponential
from PyPDF2 import PdfReader
import re
from dotenv import load_dotenv
import chardet
import pickle
from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from flask_session import Session
import tempfile
import hashlib
import boto3
from botocore.exceptions import ClientError
import logging
import sys

# Configure logging
LOG_DIR = '/var/log/flask'
if not os.path.exists(LOG_DIR):
    try:
        os.makedirs(LOG_DIR, mode=0o755)
    except Exception as e:
        # Fall back to temp directory if we can't write to /var/log/flask
        LOG_DIR = os.path.join(tempfile.gettempdir(), 'flask_logs')
        if not os.path.exists(LOG_DIR):
            os.makedirs(LOG_DIR, mode=0o755)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(os.path.join(LOG_DIR, 'application.log')),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

# Log startup information
logger.info("Starting application...")
logger.info(f"Environment: {os.getenv('FLASK_ENV', 'development')}")
logger.info(f"Python version: {sys.version}")

# Load environment variables
load_dotenv()

# Create session directory if it doesn't exist
SESSION_DIR = os.path.join(tempfile.gettempdir(), 'genas_sessions')
if not os.path.exists(SESSION_DIR):
    os.makedirs(SESSION_DIR, mode=0o777)

# Initialize Flask app
app = Flask(__name__)
app.config.update(
    SECRET_KEY=os.getenv('SECRET_KEY', 'your-secret-key-here'),
    MAX_CONTENT_LENGTH=16 * 1024 * 1024,  # 16MB max file size
    SESSION_TYPE='filesystem',
    SESSION_FILE_DIR=SESSION_DIR,
    SESSION_FILE_THRESHOLD=100,
    PERMANENT_SESSION_LIFETIME=timedelta(days=1)
)

# Initialize session interface
Session(app)

# S3 Configuration
S3_BUCKET = os.getenv('S3_BUCKET_NAME', 'genas-storage-241130')
logger.info(f"Initializing with S3 bucket: {S3_BUCKET}")

try:
    s3_client = boto3.client('s3')
    # Test S3 connection
    s3_client.list_objects_v2(Bucket=S3_BUCKET, MaxKeys=1)
    logger.info("Successfully connected to S3")
except Exception as e:
    logger.error(f"Error connecting to S3: {str(e)}")

# Ensure required directories exist
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
app.config['VERSIONS_FOLDER'] = 'versions'
app.config['DOCUMENTS_STORE'] = 'documents_store.pkl'

# Ensure required directories exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['VERSIONS_FOLDER'], exist_ok=True)

# Load existing document contents from disk
try:
    with open(app.config['DOCUMENTS_STORE'], 'rb') as f:
        document_contents = pickle.load(f)
    print(f"Loaded {len(document_contents)} documents from store")
except FileNotFoundError:
    document_contents = {}
    print("No existing document store found, starting fresh")
except Exception as e:
    print(f"Error loading document store: {e}")
    document_contents = {}

def save_documents_store():
    """Save the documents store to disk"""
    try:
        with open(app.config['DOCUMENTS_STORE'], 'wb') as f:
            pickle.dump(document_contents, f)
        print(f"Saved {len(document_contents)} documents to store")
    except Exception as e:
        print(f"Error saving document store: {e}")

# Store uploaded documents content
document_contents = {}

def detect_encoding(file_path):
    """Detect the encoding of a file using chardet"""
    with open(file_path, 'rb') as file:
        raw_data = file.read()
        result = chardet.detect(raw_data)
        return result['encoding'] if result['confidence'] > 0.7 else None

def extract_text_from_file(file):
    """Extract text from various file types"""
    if isinstance(file, str):
        file_path = file
    else:
        # If it's a FileStorage object, read it directly
        content = file.read()
        file.seek(0)  # Reset the file pointer
        
        # Get the file extension
        _, ext = os.path.splitext(file.filename)
        ext = ext.lower()
        
        if ext == '.pdf':
            try:
                reader = PdfReader(file)
                text = ''
                for page in reader.pages:
                    text += page.extract_text() + '\n'
                return text
            except Exception as e:
                logger.error(f"Error extracting PDF text: {str(e)}")
                return None
        elif ext == '.docx':
            try:
                # Save temporarily and process
                with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as temp_file:
                    temp_file.write(content)
                    temp_path = temp_file.name
                
                doc = Document(temp_path)
                text = []
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text.append(paragraph.text)
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text.append(cell.text)
                
                # Clean up
                os.unlink(temp_path)
                return '\n'.join(text)
            except Exception as e:
                logger.error(f"Error extracting DOCX text: {str(e)}")
                return None
        else:
            try:
                # Try to decode the content
                try:
                    return content.decode('utf-8')
                except UnicodeDecodeError:
                    # If UTF-8 fails, try to detect encoding
                    result = chardet.detect(content)
                    if result['encoding']:
                        return content.decode(result['encoding'], errors='replace')
                    return content.decode('utf-8', errors='replace')
            except Exception as e:
                logger.error(f"Error reading file: {str(e)}")
                return None

# Define the sections of the Acquisition Strategy document
ACQUISITION_SECTIONS = [
    "Executive Summary",
    "Program Overview",
    "Risk Assessment",
    "Market Research",
    "Competition Strategy",
    "Source Selection Planning",
    "Business Considerations",
    "Multi-Year Procurement",
    "Lease-Purchase Analysis",
    "Source of Support",
    "Environmental Considerations",
    "Security Considerations",
    "Make or Buy Program",
    "Contract Types",
    "Sustainment Strategy",
    "Supporting Documents",
    "Scratch Pad"
]

@retry(
    stop=stop_after_attempt(int(os.getenv('ANTHROPIC_MAX_RETRIES', 3))),
    wait=wait_exponential(
        multiplier=float(os.getenv('ANTHROPIC_BACKOFF_FACTOR', 2)),
        min=4,
        max=60
    ),
    retry_error_callback=lambda retry_state: {'error': f'API temporarily unavailable after {retry_state.attempt_number} attempts. Please try again in a few minutes.'}
)
def get_claude_response(system_prompt, user_prompt):
    """Get response from Claude API with retry logic"""
    try:
        print("\n=== Starting Claude API request ===")
        print(f"System prompt length: {len(system_prompt)}")
        print(f"User prompt length: {len(user_prompt)}")
        print(f"Retry settings: max_retries={os.getenv('ANTHROPIC_MAX_RETRIES', 3)}, backoff_factor={os.getenv('ANTHROPIC_BACKOFF_FACTOR', 2)}")
        
        # Debug: Print API key status (safely)
        api_key = os.getenv('ANTHROPIC_API_KEY')
        if not api_key:
            print("ERROR: No API key found!")
            raise Exception("Anthropic API key not found in environment variables")
        else:
            print(f"API key verification: {api_key[:8]}... (length: {len(api_key)})")
            
        try:
            print("\nInitializing Anthropic client...")
            client = anthropic.Client(api_key=api_key)
            
            print("Preparing message for Claude...")
            print(f"Total prompt length: {len(system_prompt) + len(user_prompt)}")
            
            print("\nSending request to Claude API...")
            message = client.messages.create(
                model="claude-3-opus-20240229",
                max_tokens=2000,
                temperature=0.7,
                system=system_prompt,
                messages=[
                    {
                        "role": "user",
                        "content": user_prompt
                    }
                ]
            )
            
            print("\nProcessing Claude response...")
            if not message:
                print("ERROR: Received empty message object")
                raise Exception("Empty response from Claude API")
            
            if not message.content:
                print("ERROR: Message object has no content")
                raise Exception("No content in Claude API response")
                
            # Extract text from the first content item
            response_text = message.content[0].text if message.content else ""
            print(f"Response received successfully (length: {len(response_text)})")
            
            print("=== Claude API request completed successfully ===\n")
            return response_text
            
        except anthropic.APIError as api_e:
            print("\n!!! Claude API Error !!!")
            print(f"Error message: {str(api_e)}")
            print(f"Error type: {type(api_e)}")
            
            if "rate_limit" in str(api_e).lower():
                raise Exception("Rate limit exceeded. Please try again in a few minutes.")
            if "invalid_api_key" in str(api_e).lower() or "unauthorized" in str(api_e).lower():
                raise Exception("Invalid API key. Please check your Anthropic API key configuration.")
                
            print("=== Claude API request failed ===\n")
            raise Exception(f"API Error: {str(api_e)}")
            
    except Exception as e:
        print("\n!!! General Error in get_claude_response !!!")
        print(f"Error message: {str(e)}")
        print(f"Error type: {type(e)}")
        raise e

CORS(app)  # Enable CORS for all routes

# Default prompt for AI interactions
DEFAULT_PROMPT = """You are an AI assistant specialized in DoD Acquisition Strategy documents. Your role is to:
1. Analyze the provided documents thoroughly
2. Answer questions based on the document content first, then supplement with your knowledge
3. Always cite which document you're referencing in your response
4. If relevant, suggest appropriate sections for including the information in an Acquisition Strategy

Please provide clear, concise responses that help create or improve DoD Acquisition Strategy documents."""

# System prompt for all interactions
SYSTEM_PROMPT = """You are an AI assistant specialized in DoD Acquisition Strategy documents. Your role is to:
1. Analyze the provided documents thoroughly
2. Answer questions based on the document content first, then supplement with your knowledge
3. Always cite which document you're referencing in your response
4. If relevant, suggest appropriate sections for including the information in an Acquisition Strategy"""

@app.route('/')
def index():
    return render_template('index.html', sections=ACQUISITION_SECTIONS)

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
        
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
        
    if file:
        try:
            logger.info(f"Processing upload for file: {file.filename}")
            
            # Check file size before processing
            file.seek(0, os.SEEK_END)
            size = file.tell()
            file.seek(0)
            
            logger.info(f"File size: {size / 1024:.1f}KB")
            
            if size > app.config['MAX_CONTENT_LENGTH']:
                return jsonify({
                    'error': f'File too large. Maximum size is {app.config["MAX_CONTENT_LENGTH"] // (1024 * 1024)}MB'
                }), 413

            filename = secure_filename(file.filename)
            
            # Extract text content directly from the file object
            logger.info("Extracting text content")
            content = extract_text_from_file(file)
            
            if content is None:
                logger.error("Failed to extract text from file")
                return jsonify({'error': 'Could not extract text from file'}), 400
            
            # Store content in S3
            logger.info("Storing content in S3")
            content_hash = store_document_content(content, filename)
            
            # Store only the reference in session
            if 'document_refs' not in session:
                session['document_refs'] = {}
            
            session['document_refs'][filename] = content_hash
            session.modified = True
            
            logger.info(f"Successfully processed file {filename}")
            
            return jsonify({
                'message': 'File uploaded and processed successfully',
                'filename': filename
            })
            
        except Exception as e:
            logger.error(f"Upload error: {str(e)}")
            return jsonify({'error': f'Error processing file: {str(e)}'}), 500
            
    return jsonify({'error': 'Invalid file'}), 400

@app.route('/save_version', methods=['POST'])
def save_version():
    """Save current document state as a version"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400

        version_name = data.get('version')
        sections = data.get('sections', [])
        
        if not version_name:
            return jsonify({'error': 'Version name is required'}), 400
            
        # Ensure version name has .json extension
        if not version_name.endswith('.json'):
            version_name += '.json'
            
        # Create versions directory if it doesn't exist
        versions_dir = app.config['VERSIONS_FOLDER']
        os.makedirs(versions_dir, exist_ok=True)
        
        version_path = os.path.join(versions_dir, secure_filename(version_name))
        
        # Save version data
        version_data = {
            'version': version_name,
            'timestamp': datetime.now().isoformat(),
            'sections': sections
        }
        
        app.logger.info(f'Saving version data: {version_data}')  # Debug log
        
        with open(version_path, 'w') as f:
            json.dump(version_data, f, indent=2)
        
        return jsonify({'message': 'Version saved successfully', 'version': version_name})
    except Exception as e:
        app.logger.error(f'Error saving version: {str(e)}')
        return jsonify({'error': str(e)}), 500

@app.route('/load_version/<version>', methods=['GET'])
def load_version(version):
    """Load a specific version"""
    try:
        version_path = os.path.join(app.config['VERSIONS_FOLDER'], secure_filename(version))
        if not os.path.exists(version_path):
            return jsonify({'error': 'Version not found'}), 404
            
        with open(version_path, 'r') as f:
            version_data = json.load(f)
            
        # Add sections array if it doesn't exist (for backward compatibility)
        if 'sections' not in version_data:
            app.logger.warning(f'Converting old version format for {version}')
            # Try to handle old format
            if 'content' in version_data:
                sections = []
                for section_data in version_data['content']:
                    sections.append({
                        'title': section_data.get('title', 'Unknown Section'),
                        'content': section_data.get('content', '')
                    })
                version_data['sections'] = sections
            else:
                return jsonify({'error': 'Invalid version file format'}), 400
            
        app.logger.info(f'Loading version data: {version_data}')  # Debug log
        return jsonify(version_data)
    except Exception as e:
        app.logger.error(f'Error loading version {version}: {str(e)}')
        return jsonify({'error': str(e)}), 500

@app.route('/delete_version/<version>', methods=['DELETE'])
def delete_version(version):
    """Delete a specific version"""
    try:
        version_path = os.path.join(app.config['VERSIONS_FOLDER'], secure_filename(version))
        if not os.path.exists(version_path):
            return jsonify({'error': 'Version not found'}), 404
            
        os.remove(version_path)
        app.logger.info(f'Deleted version: {version}')
        return jsonify({'message': f'Version {version} deleted successfully'})
    except Exception as e:
        app.logger.error(f'Error deleting version {version}: {str(e)}')
        return jsonify({'error': str(e)}), 500

@app.route('/list_versions', methods=['GET'])
def list_versions():
    """List all saved versions"""
    try:
        versions_dir = app.config['VERSIONS_FOLDER']
        if not os.path.exists(versions_dir):
            return jsonify({'versions': []})
            
        versions = []
        for filename in os.listdir(versions_dir):
            if filename.endswith('.json'):
                file_path = os.path.join(versions_dir, filename)
                try:
                    with open(file_path, 'r') as f:
                        data = json.load(f)
                        # Ensure the version has the correct format
                        if 'sections' in data or 'content' in data:
                            versions.append(filename)
                except Exception as e:
                    app.logger.warning(f'Error reading version file {filename}: {str(e)}')
                    continue
        
        versions.sort(key=lambda x: os.path.getmtime(os.path.join(versions_dir, x)), reverse=True)
        return jsonify({'versions': versions})
    except Exception as e:
        app.logger.error(f'Error listing versions: {str(e)}')
        return jsonify({'error': str(e)}), 500

@app.route('/get_versions', methods=['GET'])
def get_versions():
    versions = []
    if os.path.exists(app.config['VERSIONS_FOLDER']):
        for filename in os.listdir(app.config['VERSIONS_FOLDER']):
            if filename.endswith('.json'):
                file_path = os.path.join(app.config['VERSIONS_FOLDER'], filename)
                with open(file_path) as f:
                    version_data = json.load(f)
                    versions.append({
                        'name': version_data['version'],
                        'timestamp': version_data['timestamp']
                    })
    
    return jsonify({'versions': sorted(versions, key=lambda x: x['timestamp'], reverse=True)})

@app.route('/get_version/<version_name>', methods=['GET'])
def get_version(version_name):
    version_path = os.path.join(app.config['VERSIONS_FOLDER'], f'{version_name}.json')
    if os.path.exists(version_path):
        with open(version_path) as f:
            return jsonify(json.load(f))
    return jsonify({'error': 'Version not found'}), 404

@app.route('/documents', methods=['GET'])
def list_documents():
    """List all uploaded documents"""
    return jsonify({
        'documents': list(document_contents.keys()),
        'count': len(document_contents)
    })

@app.route('/delete_document/<filename>', methods=['DELETE'])
def delete_document(filename):
    """Delete a document from the store"""
    try:
        if filename in document_contents:
            # Remove from document store
            del document_contents[filename]
            save_documents_store()
            
            # Remove physical file if it exists
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            if os.path.exists(file_path):
                os.remove(file_path)
            
            return jsonify({'message': f'Document {filename} deleted successfully'})
        return jsonify({'error': 'Document not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/get_required_documents', methods=['POST'])
def get_required_documents():
    """Get required supporting documents based on current content"""
    try:
        # Return a simple list of required documents for now
        # This can be enhanced with AI-based document suggestions later
        required_docs = [
            {"name": "Systems Engineering Plan (SEP)", "applicability": 85},
            {"name": "Test and Evaluation Master Plan (TEMP)", "applicability": 75},
            {"name": "Cybersecurity Strategy", "applicability": 90},
            {"name": "Life Cycle Sustainment Plan (LCSP)", "applicability": 80}
        ]
        
        return jsonify({"documents": required_docs})
    except Exception as e:
        print(f"Error getting required documents: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/chat', methods=['POST'])
def chat():
    try:
        data = request.get_json()
        
        if not data or 'message' not in data:
            return jsonify({'error': 'No message provided'}), 400
            
        message = data['message']
        
        # Check if we have any documents loaded
        if 'document_refs' not in session or not session['document_refs']:
            return jsonify({'error': 'No documents loaded. Please upload documents first.'}), 400
            
        # Prepare system message with loaded documents
        documents_context = []
        failed_docs = []
        
        for filename, content_hash in session['document_refs'].items():
            content = get_document_content(content_hash, filename)
            if content:
                documents_context.append(f"Document: {filename}\nContent: {content}")
            else:
                failed_docs.append(filename)
        
        if failed_docs:
            print(f"Failed to retrieve these documents: {', '.join(failed_docs)}")
            
        if not documents_context:
            return jsonify({'error': 'Could not retrieve document contents. Please try uploading again.'}), 500
            
        documents_text = "\n\n".join(documents_context)
        
        # Initialize retry parameters
        max_retries = 3
        base_delay = 5
        
        for attempt in range(max_retries):
            try:
                messages = [
                    {
                        "role": "system",
                        "content": f"You are a helpful AI assistant. Here are the documents to reference:\n\n{documents_text}"
                    },
                    {
                        "role": "user",
                        "content": message
                    }
                ]
                
                if 'conversation_history' in session:
                    messages[1:1] = session['conversation_history']
                
                response = client.messages.create(
                    model=os.getenv('MODEL_NAME', "claude-2"),
                    max_tokens=int(os.getenv('MAX_TOKENS', 100000)),
                    temperature=float(os.getenv('TEMPERATURE', 0.7)),
                    messages=messages
                )
                
                if 'conversation_history' not in session:
                    session['conversation_history'] = []
                    
                session['conversation_history'].extend([
                    {"role": "user", "content": message},
                    {"role": "assistant", "content": response.content[0].text}
                ])
                
                if len(session['conversation_history']) > 10:
                    session['conversation_history'] = session['conversation_history'][-10:]
                    
                session.modified = True
                
                return jsonify({'response': response.content[0].text})
                
            except anthropic.RateLimitError as e:
                if attempt == max_retries - 1:
                    print(f"Rate limit error after {max_retries} attempts: {str(e)}")
                    return jsonify({
                        'error': 'Rate limit exceeded. Please try again in about an hour.',
                        'retry_after': '3600'
                    }), 429
                    
                delay = base_delay * (2 ** attempt)
                print(f"Rate limit hit, attempt {attempt + 1}/{max_retries}. Waiting {delay} seconds...")
                time.sleep(delay)
                continue
                
            except Exception as e:
                print(f"Error in chat: {str(e)}")
                return jsonify({'error': f'An error occurred: {str(e)}'}), 500
                
    except Exception as e:
        print(f"Error processing request: {str(e)}")
        return jsonify({'error': f'An error occurred: {str(e)}'}), 500

@app.route('/print_document', methods=['POST'])
def print_document():
    try:
        # Get the document content from the request
        data = request.get_json()
        if not data or 'sections' not in data:
            return jsonify({'error': 'No document content provided'}), 400

        sections = data['sections']
        
        # Create a BytesIO buffer for the PDF
        buffer = BytesIO()
        
        # Create the PDF document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )

        # Create styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=20,
            spaceAfter=10,
            textColor=colors.HexColor('#2c3e50')
        )
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceBefore=6,
            spaceAfter=12,
            leading=14
        )
        date_style = ParagraphStyle(
            'CustomDate',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.gray,
            alignment=1  # Center alignment
        )

        # Build the document content
        story = []
        
        # Add title
        story.append(Paragraph("Acquisition Strategy Document", title_style))
        story.append(Paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y')}", date_style))
        story.append(Spacer(1, 30))

        # Add sections
        for title, content in sections.items():
            if content.strip():
                story.append(Paragraph(title, heading_style))
                # Split content into paragraphs
                paragraphs = content.split('\n')
                for para in paragraphs:
                    if para.strip():
                        story.append(Paragraph(para.strip(), normal_style))
                story.append(Spacer(1, 12))

        # Build the PDF
        doc.build(story)

        # Reset buffer position
        buffer.seek(0)
        
        # Create response
        response = send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name='acquisition_strategy.pdf'
        )

        return response

    except Exception as e:
        print(f"Error generating PDF: {str(e)}")
        return jsonify({'error': 'Error generating PDF'}), 500

def get_or_create_document():
    doc_path = os.path.join(app.config['UPLOAD_FOLDER'], 'current_document.json')
    
    if not os.path.exists(doc_path):
        # Create default document structure
        default_document = {
            "title": "Acquisition Strategy",
            "sections": [
                {
                    "title": section,
                    "content": "",
                    "subsections": []
                }
                for section in ACQUISITION_SECTIONS
            ]
        }
        
        # Ensure upload directory exists
        os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
        
        # Save default document
        with open(doc_path, 'w') as f:
            json.dump(default_document, f, indent=4)
        
        return default_document
    
    # Load existing document
    with open(doc_path, 'r') as f:
        return json.load(f)

@app.route('/update_section', methods=['POST'])
def update_section():
    try:
        data = request.get_json()
        section = data.get('section')
        content = data.get('content')
        
        if not section or not content:
            return jsonify({'error': 'Missing section or content'}), 400
            
        # Get or create document
        document = get_or_create_document()
            
        # Update the specified section
        section_found = False
        for section_data in document['sections']:
            if section_data['title'] == section:
                section_data['content'] = content
                section_found = True
                break
                
        if not section_found:
            return jsonify({'error': f'Section {section} not found'}), 404
            
        # Save the updated document
        doc_path = os.path.join(app.config['UPLOAD_FOLDER'], 'current_document.json')
        with open(doc_path, 'w') as f:
            json.dump(document, f, indent=4)
            
        return jsonify({'message': f'Successfully updated {section}'}), 200
        
    except Exception as e:
        app.logger.error(f'Error updating section: {str(e)}')
        return jsonify({'error': 'Internal server error'}), 500

@app.route('/get_document', methods=['GET'])
def get_document():
    try:
        document = get_or_create_document()
        return jsonify(document)
    except Exception as e:
        app.logger.error(f'Error getting document: {str(e)}')
        return jsonify({'error': 'Internal server error'}), 500

@app.route('/clear_session', methods=['POST'])
def clear_session():
    """Clear the current session state"""
    try:
        print("Clearing session state...")
        
        # Clear document contents
        global document_contents
        document_contents = {}
        save_documents_store()
        
        # Clear uploaded files
        upload_dir = app.config['UPLOAD_FOLDER']
        for filename in os.listdir(upload_dir):
            file_path = os.path.join(upload_dir, filename)
            try:
                if os.path.isfile(file_path):
                    os.unlink(file_path)
                    print(f"Deleted file: {filename}")
            except Exception as e:
                print(f"Error deleting {filename}: {e}")
        
        print("Session cleared successfully")
        return jsonify({'message': 'Session cleared successfully'})
        
    except Exception as e:
        print(f"Error clearing session: {e}")
        return jsonify({'error': str(e)}), 500

def load_prompts():
    """Load prompts from JSON file"""
    try:
        with open('static/prompts.json', 'r') as f:
            prompts = json.load(f)
            print(f"Loaded {len(prompts)} prompts from prompts.json")
            return prompts
    except Exception as e:
        print(f"Error loading prompts: {e}")
        return {
            'default': {
                'id': 'default',
                'name': 'Default Prompt',
                'description': 'Default prompt for DoD Acquisition Strategy assistance',
                'prompt': DEFAULT_PROMPT
            }
        }

@app.route('/prompts', methods=['GET'])
def get_prompts():
    """Return all available prompts"""
    try:
        prompts = load_prompts()
        return jsonify(prompts)
    except Exception as e:
        print(f"Error in get_prompts: {e}")
        return jsonify({
            'prompts': [{
                'id': 'default',
                'name': 'Default Prompt',
                'description': 'Default prompt for DoD Acquisition Strategy assistance',
                'prompt': DEFAULT_PROMPT,
                'category': 'General'
            }]
        })

@app.route('/prompt/<prompt_id>', methods=['GET'])
def get_prompt(prompt_id):
    """Return a specific prompt by ID"""
    try:
        prompts = load_prompts()
        if 'prompts' in prompts:
            for prompt in prompts['prompts']:
                if prompt['id'] == prompt_id:
                    return jsonify(prompt)
        return jsonify({'error': 'Prompt not found'}), 404
    except Exception as e:
        print(f"Error in get_prompt: {e}")
        if prompt_id == 'default':
            return jsonify({
                'id': 'default',
                'name': 'Default Prompt',
                'description': 'Default prompt for DoD Acquisition Strategy assistance',
                'prompt': DEFAULT_PROMPT
            })
        return jsonify({'error': str(e)}), 500

def store_document_content(content, filename):
    """Store document content in S3 and return content hash."""
    try:
        # Generate a unique hash for the content
        content_hash = hashlib.md5(content.encode()).hexdigest()
        
        # Store the content in S3
        s3_key = f'documents/{content_hash}/{filename}'
        s3_client.put_object(
            Bucket=S3_BUCKET,
            Key=s3_key,
            Body=content.encode(),
            ContentType='text/plain'
        )
        logger.info(f"Stored document in S3: {s3_key}")
        return content_hash
    except Exception as e:
        logger.error(f"Error storing document in S3: {str(e)}")
        raise

def get_document_content(content_hash, filename):
    """Retrieve document content from S3."""
    try:
        s3_key = f'documents/{content_hash}/{filename}'
        response = s3_client.get_object(Bucket=S3_BUCKET, Key=s3_key)
        content = response['Body'].read().decode()
        return content
    except ClientError as e:
        if e.response['Error']['Code'] == 'NoSuchKey':
            logger.error(f"Document not found in S3: {s3_key}")
            return None
        raise
    except Exception as e:
        logger.error(f"Error retrieving document from S3: {str(e)}")
        raise

@app.route('/health')
def health_check():
    return jsonify({"status": "healthy"}), 200

if __name__ == '__main__':
    # Only use debug mode when running locally
    is_dev = os.getenv('FLASK_ENV') == 'development'
    port = int(os.getenv('PORT', 8000))
    app.run(host='0.0.0.0', port=port, debug=is_dev)
